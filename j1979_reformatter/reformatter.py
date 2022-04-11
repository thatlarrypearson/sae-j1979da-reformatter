# sae-j1979da-reformatter/j1979_reformatter/reformatter.py
"""
Reformat Excel OBD Interface Standard SAE-J1979 to Word
"""
from sys import stdout, stderr
from argparse import ArgumentParser
from typing import Tuple
from re import sub as resub
from pprint import PrettyPrinter
from openpyxl import load_workbook
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from obd.commands import __mode1__, __mode9__
from telemetry_obd.add_commands import NEW_COMMANDS

ALL_COMMANDS = __mode1__ + __mode9__ + NEW_COMMANDS

DEFAULT_WORD = "SAE-J1979DA.docx"
DEFAULT_XLSX = "SAE J1979DA_202104.xlsx"
ANNEX_B = "Annex B - Parameter IDs"
ANNEX_G = "Annex G - InfoType IDs"

DOCUMENT_TITLE="SAE J1979 Standard Extract"

WHITE_SPACE=r"[\t\r\n]+"

def command_line_options()->dict:
    parser = ArgumentParser(prog="reformatter", description="Reformat Excel OBD Interface Standard SAE-J1979 to Word")

    parser.add_argument(
        "--commands",
        help="""
            Command name list to include in Word report generation.
            Command names come from 'telemetry-obd' (including 'python-OBD') package.
            Comma separated list.  e.g. "SPEED,RPM,FUEL_RATE".
        """,
    )

    parser.add_argument(
        "--annex_b",
        help="""
            Annex B Parameter IDs (PID) list to include in Word report generation.
            Comma separated list. e.g. "0x01,0x0F,0xAF".
        """,
    )

    parser.add_argument(
        "--annex_g",
        help="""
            Annex G Info Type IDs (PID) list to include in Word report generation.
            Comma separated list. e.g. "0x01,0x0F,0xAF".
        """,
    )

    parser.add_argument(
        "--word",
        help=f"""Word output file.
                File can be either a full or relative path name.
                If the file already exists, it will be overwritten.
                Defaults to '{DEFAULT_WORD}'.
                """,
        default=DEFAULT_WORD,
    )

    parser.add_argument(
        "--xlsx",
        help=f"""
            Excel version of SAE standard document J1979DA file name.
            Defaults to "{DEFAULT_XLSX}".
        """,
        default=DEFAULT_XLSX,
    )

    parser.add_argument(
        "--verbose",
        help="Turn verbose output on. Default is off.",
        default=False,
        action='store_true'
    )

    return vars(parser.parse_args())

def get_annex_pid_from_command(command:str, verbose=False)->Tuple[str, str]:
    """
    Get command info from __mode9__, __mode1__ or NEW_COMMANDS
    """
    for cmd in __mode1__:
        if cmd.name == command:
            if verbose:
                print(f"get_annex_pid_from_command: mode1: {command}", file=stderr)
            return 'annex_b', f"0x{(cmd.command.decode('utf-8'))[2:4]}"
    for cmd in __mode9__:
        if cmd.name == command:
            if verbose:
                print(f"get_annex_pid_from_command: mode9: {command}", file=stderr)
            return 'annex_g', f"0x{(cmd.command.decode('utf-8'))[2:4]}"
    for cmd in NEW_COMMANDS:
        if cmd.name == command:
            mode = (cmd.command.decode('utf-8'))[:2]
            if mode == '01':
                if verbose:
                    print(f"get_annex_pid_from_command: NEW_COMMANDS: mode1: {command}", file=stderr)
                return 'annex_b', f"0x{(cmd.command.decode('utf-8'))[2:4]}"
            if mode == '09':
                if verbose:
                    print(f"get_annex_pid_from_command: NEW_COMMANDS: mode9: {command}", file=stderr)
                return 'annex_g', f"0x{(cmd.command.decode('utf-8'))[2:4]}"

    if verbose:
        print(f"get_annex_pid_from_command: command {command} not found", file=stderr)

    return None, None

def find_command_name_for_pid(annex, pid, verbose=False)->str:
    """
    Map PID to command name starting with command names. 
    """
    for cmd in ALL_COMMANDS:
        cmd_mode = (cmd.command.decode('utf-8'))[:2]
        cmd_pid = (cmd.command.decode('utf-8'))[2:4]

        if annex == 'annex_b' and cmd_mode == '01' and pid[2:4] == cmd_pid:
            if verbose:
                print(f"find_command_name_for_pid({annex}, {pid[2:4]}): {cmd.name}", file=stderr)
            return cmd.name
        if annex == 'annex_g' and cmd_mode == '09' and pid[2:4] == cmd_pid:
            if verbose:
                print(f"find_command_name_for_pid({annex}, {pid[2:4]}): {cmd.name}", file=stderr)
            return cmd.name

    if verbose:
        print(f"find_command_name_for_pid({annex}, {pid}): None", file=stderr)

    return None

def get_annex_lists(commands:list, annex_b_pids:list, annex_g_pids:list, verbose=False)->Tuple[dict, dict]:
    """
    Return Annex B and Annex G list dictionaries
    """
    annex_b_list = {}
    for pid in annex_b_pids:
        annex_b_list[pid] = {
            'name': find_command_name_for_pid('annex_b', pid, verbose=verbose),
            'mode': '01',
            'pid': pid,
            'annex': 'annex_b',
        }

    annex_g_list = {}
    for pid in annex_g_pids:
        annex_g_list[pid] = {
            'name': find_command_name_for_pid('annex_g', pid, verbose=verbose),
            'mode': '09',
            'pid': pid,
            'annex': 'annex_g',
        }

    for cmd in commands:
        annex, pid = get_annex_pid_from_command(cmd, verbose=verbose)
        if pid is None:
            if verbose:
                print(f"get_annex_lists: command {cmd}: pid is None", file=stderr)
            continue
        if annex == 'annex_b':
            annex_b_list[pid] = {
                'name': cmd,
                'mode': '01',
                'pid': pid,
                'annex': 'annex_b',
            }
        elif annex == 'annex_g':
            annex_g_list[pid] = {
                'name': cmd,
                'mode': '09',
                'pid': pid,
                'annex': 'annex_g',
            }

    # if verbose:
    #     print(f"get_annex_lists: annex_b_list {annex_b_list}", file=stderr)
    #     print(f"get_annex_lists: annex_g_list {annex_g_list}", file=stderr)

    return annex_b_list, annex_g_list

def get_sheet_header(xlsx, sheet_name:str, verbose=False) -> list:
    """
    Return headers from a sheet as a list of strings
    """
    ws = xlsx[sheet_name]
    header =[]
    for i, row in enumerate(ws.iter_rows(min_row=1, min_col=1, max_row=1), start=1):
        for j, cell in enumerate(row, start=1):
            if cell.value:
                new_value = resub(WHITE_SPACE, "", cell.value)
                header.append(new_value)

    # if verbose:
    #     print(f"get_sheet_header: header: {header}", file=stderr)

    return header

def sheet_pid_search(ws, pid:str, info:dict, verbose=False)->Tuple [int,int]:
    """
    Get the first and last row where pid (e.g. 0xA0) is found 
    """
#     for row in enumerate(ws.iter_rows(min_row=1, max_col=1, max_row=1), start=1):
#        for j, cell in enumerate(row, start=1):
    first_row_found = False
    for i, row in enumerate(ws.iter_rows(min_row=1, min_col=1, max_col=1), start=1):
        for j, cell in enumerate(row, start=1):
            if not first_row_found and isinstance(cell.value, str) and pid == cell.value:
                info['first_row'] = i
                first_row_found = True
                continue
            elif first_row_found and isinstance(cell.value, str) and cell.value.startswith("0x") and pid != cell.value:
                info['last_row'] = i - 1
                return info['first_row'], info['last_row']

    return None, None

def get_pid_header(ws, annex_header:list, pid:str, info:dict, verbose=False)->list:
    pid_header = {}
    for row in ws.iter_rows(min_row=info['first_row'], min_col=1, max_row=info['first_row']):
        for j, cell in enumerate(row, start=1):
            if cell.value:
                pid_header[annex_header[j-1]] = cell.value

    return pid_header

def get_pid_fields(ws, annex_header:list, pid:str, info:dict, verbose=False)->dict:
    # sourcery skip: dict-comprehension, inline-immediately-returned-variable
    pid_fields = {}
    for i, row in enumerate(ws.iter_rows(min_row=(info['first_row'] + 1), min_col=1, max_row=info['last_row']), start=1):
        pid_fields[i] = {}
        for j, cell in enumerate(row, start=1):
            if cell.value:
                pid_fields[i][annex_header[j-1]] = cell.value

    return pid_fields

def sheet(xlsx, sheet_name:str, annex_header:list, annex_items:dict, verbose=False)->dict:
    """
    for each annex item or pid, add data from annex sheet into dictionary
    """
    ws = xlsx[sheet_name]
    for key, value in annex_items.items():
        first_row, last_row = sheet_pid_search(ws, key, value)
        # if verbose:
        #     print(f"sheet: {value['annex']}/{key}/{value['name']}: first_row {first_row}, last_row {last_row}", file=stderr)
        value['pid_header'] = get_pid_header(ws, annex_header, key, value, verbose=verbose)
        # if verbose:
        #     print(f"sheet: {value['annex']}/{key}/{value['name']}: pid_header: {value['pid_header']}", file=stderr)
        value['pid_fields'] = get_pid_fields(ws, annex_header, key, value)
        # if verbose:
        #     print(f"sheet: {value['annex']}/{key}/{value['name']}: pid_fields: {value['pid_fields']}", file=stderr)

    return annex_items


def spreadsheet(name:str, commands:list, annex_b_pids:list, annex_g_pids:list, verbose=False)->Tuple [list, dict, list, dict]:
    """
    Load an Excel Spreadsheet and turn it into a dictionary with each of the
    commands/PIDs 
    """
    xlsx = load_workbook(name)

    if ANNEX_B not in xlsx.sheetnames:
        raise ValueError(f"Bad Spreadsheet {name}: sheet {ANNEX_B} not in {xls.sheetnames}")

    if ANNEX_G not in xlsx.sheetnames:
        raise ValueError(f"Bad Spreadsheet {name}: sheet {ANNEX_G} not in {xls.sheetnames}")

    annex_b_items, annex_g_items = get_annex_lists(commands, annex_b_pids, annex_g_pids, verbose=verbose)

    annex_b_header = get_sheet_header(xlsx, ANNEX_B, verbose=verbose)
    annex_g_header = get_sheet_header(xlsx, ANNEX_G, verbose=verbose)

    sheet(xlsx, ANNEX_B, annex_b_header, annex_b_items, verbose=verbose)
    sheet(xlsx, ANNEX_G, annex_g_header, annex_g_items, verbose=verbose)

    xlsx.close()

    return annex_b_header, annex_b_items, annex_g_header, annex_g_items

# Adapted from https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)


def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def create_word_document(annex_b_headers:dict, annex_b_items:dict, annex_g_headers:dict, annex_g_items:dict, file_name=DEFAULT_WORD, verbose=False):
    """
    Create word document from Excel workbook data.
    """
    document = Document()
    # document.add_heading(DOCUMENT_TITLE, 0)

    # header
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = f"\t\t{DOCUMENT_TITLE}: {file_name}"

    # footer
    # footer = section.footer
    # paragraph.text = "Left Text\tCenter Text\tRight Text"
    # paragraph.style = document.styles["Header"]
    add_page_number(document.sections[0].footer.paragraphs[0].add_run())

    # for each Annex B command
    for pid, info in annex_b_items.items():
        # annex b command
        if verbose:
            print(f"create_word_document: 0x{info['mode']}, {pid}", file=stderr)
        document.add_heading(f"{ANNEX_B}: {pid}: {info['name']} ({info['pid_header']['Description']})", 2)
        if 'Comment' in info['pid_header']:
            document.add_paragraph(info['pid_header']['Comment'])
        for i, field in info['pid_fields'].items():
            if 'Data Byte' in field:
                document.add_heading(f"{i}. Data Byte {field['Data Byte']}, {field['Description']}", 3)
                if 'Comment' in field:
                    document.add_paragraph(field['Comment'])
                if 'US OBD Regulatory term used' in field:
                    document.add_heading("US OBD Regulatory term used", 4)
                    document.add_paragraph(resub(WHITE_SPACE, " ", field['US OBD Regulatory term used']))
                table = document.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Data Byte'
                hdr_cells[1].text = 'Maximum Value'
                hdr_cells[2].text = 'Minimum Value'
                hdr_cells[3].text = 'Scaling/Bit'
                row_cells = table.add_row().cells
                row_cells[0].text = str(field['Data Byte'])
                if 'Max. Value' in field:
                    row_cells[1].text = str(field['Max. Value'])
                if 'Min. Value' in field:
                    row_cells[2].text = str(field['Min. Value'])
                if 'Scaling/bit' in field:
                    row_cells[3].text = str(field['Scaling/bit'])
        document.add_page_break()

    # for each Annex G command
    for pid, info in annex_g_items.items():
        # annex g command
        if verbose:
            print(f"create_word_document: 0x{info['mode']}, {pid}", file=stderr)
        document.add_heading(f"{ANNEX_G}: {pid}: {info['name']} ({info['pid_header']['Description']})", 2)
        if 'Comment' in info['pid_header']:
            document.add_paragraph(info['pid_header']['Comment'])
        for i, field in info['pid_fields'].items():
            if 'Data Byte' in field:
                document.add_heading(f"{i}. Data Byte {field['Data Byte']}, {field['Description']}", 3)
                if 'Comment' in field:
                    document.add_paragraph(field['Comment'])
                if 'US OBD Regulatory term used' in field:
                    document.add_heading("US OBD Regulatory term used", 4)
                    document.add_paragraph(resub(WHITE_SPACE, " ", field['US OBD Regulatory term used']))
                table = document.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Data Byte'
                hdr_cells[1].text = 'Maximum Value'
                hdr_cells[0].text = 'Minimum Value'
                hdr_cells[1].text = 'Scaling/Bit'
                row_cells = table.add_row().cells
                row_cells[0].text = str(field['Data Byte'])
                if 'Max. Value' in field:
                    row_cells[1].text = str(field['Max. Value'])
                if 'Min. Value' in field:
                    row_cells[2].text = str(field['Min. Value'])
                if 'Scaling/bit' in field:
                    row_cells[3].text = str(field['Scaling/bit'])
            else:
                if 'Description' in field:
                    document.add_heading(f"{i}. {field['Description']}", 1)
                if 'Comment' in field:
                    document.add_paragraph(field['Comment'])
               
        document.add_page_break()

    document.save(file_name)

def main():  # sourcery skip: assign-if-exp, use-assigned-variable
    args = command_line_options()

    verbose = args['verbose']

    if args['commands']:
        commands = (args['commands']).split(sep=',')
    else:
        commands = []

    if args['annex_b']:
        annex_b = (args['annex_b']).split(sep=',')
    else:
        annex_b = []

    if args['annex_g']:
        annex_g = (args['annex_g']).split(sep=',')
    else:
        annex_g = []

    xlsx_file_name = args['xlsx']
    word_file_name = args['word']

    if verbose:
        print(f"verbose: {args['verbose']}", file=stderr)
        print(f"commands: {args['commands']}", file=stderr)
        print(f"annex_b: {annex_b}", file=stderr)
        print(f"annex_g: {annex_g}", file=stderr)
        print(f"xlsx: {xlsx_file_name}", file=stderr)
        print(f"word: {word_file_name}", file=stderr)

    annex_b_headers, annex_b_items, annex_g_headers, annex_g_items = spreadsheet(xlsx_file_name, commands, annex_b, annex_g, verbose=verbose)

    if verbose:
        pp = PrettyPrinter(indent=2, width=150, stream=stderr)

        print("annex_b_items:")
        pp.pprint(annex_b_items)

        print("annex_g_items:")
        pp.pprint(annex_g_items)

    create_word_document(annex_b_headers, annex_b_items, annex_g_headers, annex_g_items, file_name=word_file_name, verbose=verbose)

if __name__ == "__main__":
    main()
